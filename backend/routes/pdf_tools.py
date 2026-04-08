import io
import os
import fitz  # PyMuPDF
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from typing import Optional
from docx import Document as DocxDocument

router = APIRouter(prefix="/pdf-tools", tags=["PDF Tools"])


async def _read_pdf(file: UploadFile) -> fitz.Document:
    """Read an uploaded file into a PyMuPDF Document."""
    data = await file.read()
    return fitz.open(stream=data, filetype="pdf")


@router.post("/merge")
async def merge_pdfs(files: list[UploadFile] = File(...)):
    """Merge multiple PDF files into one."""
    if len(files) < 2:
        raise HTTPException(400, "At least 2 PDF files required for merging.")

    merged = fitz.open()
    for f in files:
        pdf = await _read_pdf(f)
        merged.insert_pdf(pdf)
        pdf.close()

    buf = io.BytesIO(merged.tobytes())
    merged.close()

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=merged.pdf"},
    )


@router.post("/split")
async def split_pdf(
    file: UploadFile = File(...),
    start_page: int = Form(1),
    end_page: int = Form(0),
):
    """Split a PDF — extract a range of pages."""
    pdf = await _read_pdf(file)
    total = len(pdf)

    if end_page <= 0:
        end_page = total
    start_page = max(1, start_page)
    end_page = min(total, end_page)

    if start_page > end_page:
        raise HTTPException(400, "start_page must be <= end_page")

    out = fitz.open()
    out.insert_pdf(pdf, from_page=start_page - 1, to_page=end_page - 1)
    pdf.close()

    buf = io.BytesIO(out.tobytes())
    out.close()

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=split_{start_page}-{end_page}.pdf"},
    )


@router.post("/compress")
async def compress_pdf(file: UploadFile = File(...)):
    """Compress a PDF to reduce file size."""
    pdf = await _read_pdf(file)

    # Deflate and clean — write compressed bytes directly
    buf = io.BytesIO(pdf.tobytes(garbage=4, deflate=True, clean=True))
    pdf.close()

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=compressed.pdf"},
    )


@router.post("/rotate")
async def rotate_pdf(
    file: UploadFile = File(...),
    angle: int = Form(90),
):
    """Rotate all pages in a PDF by the given angle (90, 180, 270)."""
    if angle not in (90, 180, 270):
        raise HTTPException(400, "Angle must be 90, 180, or 270")

    pdf = await _read_pdf(file)
    for page in pdf:
        page.set_rotation(page.rotation + angle)

    buf = io.BytesIO(pdf.tobytes())
    pdf.close()

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=rotated.pdf"},
    )


@router.post("/add-page-numbers")
async def add_page_numbers(file: UploadFile = File(...)):
    """Add page numbers to the bottom of each page."""
    pdf = await _read_pdf(file)

    for i, page in enumerate(pdf):
        rect = page.rect
        text = f"Page {i + 1} of {len(pdf)}"
        point = fitz.Point(rect.width / 2 - 30, rect.height - 20)
        page.insert_text(point, text, fontsize=10, color=(0.4, 0.4, 0.4))

    buf = io.BytesIO(pdf.tobytes())
    pdf.close()

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=numbered.pdf"},
    )


@router.post("/watermark")
async def add_watermark(
    file: UploadFile = File(...),
    text: str = Form("CONFIDENTIAL"),
):
    """Add a diagonal text watermark to every page."""
    pdf = await _read_pdf(file)

    for page in pdf:
        rect = page.rect
        # Insert diagonal watermark
        page.insert_text(
            fitz.Point(rect.width / 4, rect.height / 2),
            text,
            fontsize=48,
            color=(0.85, 0.85, 0.85),
            rotate=45,
        )

    buf = io.BytesIO(pdf.tobytes())
    pdf.close()

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=watermarked.pdf"},
    )


@router.post("/to-text")
async def pdf_to_text(file: UploadFile = File(...)):
    """Extract all text from a PDF."""
    pdf = await _read_pdf(file)
    text_parts = []
    for i, page in enumerate(pdf):
        text_parts.append(f"--- Page {i + 1} ---\n{page.get_text()}")
    pdf.close()

    content = "\n\n".join(text_parts)
    buf = io.BytesIO(content.encode("utf-8"))

    return StreamingResponse(
        buf,
        media_type="text/plain",
        headers={"Content-Disposition": "attachment; filename=extracted.txt"},
    )


@router.post("/to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    """Convert PDF to DOCX (text extraction based)."""
    pdf = await _read_pdf(file)

    doc = DocxDocument()
    doc.add_heading("Converted Document", 0)

    for i, page in enumerate(pdf):
        doc.add_heading(f"Page {i + 1}", level=1)
        text = page.get_text()
        for paragraph in text.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

    pdf.close()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=converted.docx"},
    )
